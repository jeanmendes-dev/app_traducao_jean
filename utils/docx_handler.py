"""
Traducao de documentos .docx preservando a formatacao original.

Estrategia: o texto e traduzido no nivel de "run" (o menor trecho de texto
que compartilha uma mesma formatacao dentro de um paragrafo no formato
OOXML). Como o texto de cada run e substituido no proprio objeto run
(sem remover ou recriar o run), todos os atributos de formatacao
(fonte, tamanho, negrito, italico, cor, realce, etc.) sao preservados
automaticamente pelo python-docx -- nada na arvore XML alem do conteudo
textual e tocado.

Sao percorridos: paragrafos do corpo do documento, tabelas (inclusive
tabelas aninhadas dentro de celulas) e os cabecalhos/rodapes de cada secao
(incluindo variantes de primeira pagina e pagina par, quando existirem).

Imagens, graficos, quebras de pagina, numeracao automatica de paginas
(campos de campo do Word) e a estrutura de tabelas nao sao alteradas --
apenas o texto visivel dos runs e substituido.

Limitacoes conhecidas (reportadas ao usuario na interface):
  - Caixas de texto flutuantes, WordArt e objetos incorporados (SmartArt,
    OLE) nao sao expostos pelo python-docx e portanto nao sao traduzidos.
  - Quando um mesmo paragrafo mistura formatacoes diferentes em palavras
    distintas (ex.: apenas uma palavra em negrito), a traducao e feita
    run a run para preservar exatamente essa formatacao; isso pode, em
    casos raros, afetar a fluidez gramatical de frases divididas entre
    runs com formatacao diferente.
"""

import io
import re

import docx

BATCH_SIZE = 40

_NUMERIC_ONLY_RE = re.compile(r'^[\d\s\.\,\-\/\:\(\)%]+$')


def _is_translatable(text: str) -> bool:
    if not text or not text.strip():
        return False
    # Evita traduzir trechos puramente numericos/simbolicos, que tendem a
    # ser numeros de pagina, datas ja formatadas ou codigos.
    if _NUMERIC_ONLY_RE.match(text.strip()):
        return False
    return True


def _collect_paragraph_runs(paragraphs, out):
    for p in paragraphs:
        for r in p.runs:
            out.append(r)


def _collect_table_runs(tables, out):
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                _collect_paragraph_runs(cell.paragraphs, out)
                # tabelas aninhadas dentro de uma celula
                if cell.tables:
                    _collect_table_runs(cell.tables, out)


def _collect_all_runs(document):
    runs = []
    _collect_paragraph_runs(document.paragraphs, runs)
    _collect_table_runs(document.tables, runs)

    for section in document.sections:
        header_footer_parts = [
            section.header, section.footer,
        ]
        # first_page_header/footer e even_page_header/footer so existem
        # quando a opcao correspondente esta habilitada na secao
        for attr in ("first_page_header", "first_page_footer",
                     "even_page_header", "even_page_footer"):
            part = getattr(section, attr, None)
            if part is not None:
                header_footer_parts.append(part)

        for part in header_footer_parts:
            if part is None:
                continue
            _collect_paragraph_runs(part.paragraphs, runs)
            if part.tables:
                _collect_table_runs(part.tables, runs)

    return runs


def translate_docx_bytes(file_bytes: bytes, translate_fn, progress_callback=None):
    """
    Traduz o conteudo textual de um .docx, preservando formatacao.

    `translate_fn`: funcao que recebe uma lista de strings e devolve uma
    lista de strings traduzidas, na mesma ordem.
    `progress_callback(done, total)`: chamada apos cada lote traduzido.

    Retorna os bytes do novo .docx.
    """
    document = docx.Document(io.BytesIO(file_bytes))

    all_runs = _collect_all_runs(document)
    targets = [r for r in all_runs if _is_translatable(r.text)]

    total = len(targets)
    done = 0
    if progress_callback:
        progress_callback(0, max(total, 1))

    for start in range(0, total, BATCH_SIZE):
        batch = targets[start:start + BATCH_SIZE]
        original_texts = [r.text for r in batch]
        translated_texts = translate_fn(original_texts)
        for run, new_text in zip(batch, translated_texts):
            run.text = new_text
        done += len(batch)
        if progress_callback:
            progress_callback(done, max(total, 1))

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def count_translatable_segments(file_bytes: bytes) -> int:
    document = docx.Document(io.BytesIO(file_bytes))
    runs = _collect_all_runs(document)
    return sum(1 for r in runs if _is_translatable(r.text))


def extract_sample_text(file_bytes: bytes, max_chars: int = 4000) -> str:
    """Extrai uma amostra de texto do documento, usada na deteccao de idioma."""
    document = docx.Document(io.BytesIO(file_bytes))
    parts = []
    total_len = 0
    for p in document.paragraphs:
        if p.text.strip():
            parts.append(p.text)
            total_len += len(p.text)
        if total_len >= max_chars:
            break
    return "\n".join(parts)[:max_chars]
